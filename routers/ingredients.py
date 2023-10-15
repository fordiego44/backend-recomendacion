from fastapi import APIRouter, HTTPException, status
from pydantic import BaseModel
from db.models.url import Url
import dateparser
import matplotlib.pyplot as plt
from db.client import db_client
from db.schemas.documento import document_schema, documents_schema
from db.models.documento import Documento
import os

import requests
from bs4 import BeautifulSoup
import re
from pdf2docx import Converter
from docx import Document
import pandas as pd
from statsmodels.tsa.arima.model import ARIMA
import json 
import pickle
from datetime import datetime

# import pandas as pd

router = APIRouter( 
    prefix="/ingredients",
    tags=["ingredients"],
    responses={status.HTTP_404_NOT_FOUND: {"message": "No encontrado"}}
)

@router.post("/pricesPrediction", response_model=dict)
async def pricesPrediction(data: dict): 
    # tipoPrecio = 0
    # proveedor = "Mercado Sobre Ruedas"
    # # archivoPath = "output-24%20de%20Septiembre%20de%202023.docx"
    # archivoPath = "static\\docs\\word\\output-24%20de%20Septiembre%20de%202023.docx"
    # Crea un DataFrame vacío donde concatenar los resultados.
    df = pd.DataFrame()
    documentos = data["archivosPath"]

    for documento in documentos:
        df_old = await getPriceProveedor(data["proveedor"], data["tipoPrecio"], documento) 
     
        # Verifica si el DataFrame es válido antes de concatenarlo
        if not df_old.empty:
   
            df = pd.concat([df, df_old], ignore_index=True)
       
    # Inicializa un DataFrame vacío para almacenar los resultados
    resultados_df = pd.DataFrame(columns=['ingrediente', 'precio'])
    

    # Obtiene la lista de ingredientes únicos
    ingredientes_unicos = df['ingrediente'].unique()
 
 
    # Itera sobre cada ingrediente
    for ingrediente in ingredientes_unicos:
        # Llama a tu función arima_predictor
        zan_df, zan_pred = arima_predictor(df, ingrediente, p=1, d=1, q=1)
        
        # Extrae el primer precio de las predicciones
        precio = zan_pred.iloc[0]
        
        # Crear un nuevo DataFrame con los valores a agregar
        nuevo_df = pd.DataFrame({'ingrediente': ingrediente, 'precio': precio}, index=[0])

        # Usar concat para agregar el nuevo DataFrame al existente
        resultados_df = pd.concat([resultados_df, nuevo_df], ignore_index=True)

    # Muestra el DataFrame de resultados
    # resultados_df
    # Convertir el DataFrame a JSON
    json_data = resultados_df.to_json(orient='records')

    # return data_list
    # Convertir el DataFrame a JSON y luego a una lista de diccionarios
    json_resultados = resultados_df.to_json(orient="records")
    
    # Cargar el JSON como una lista de diccionarios
    lista_resultados = json.loads(json_resultados)
    
    return {
        "ingredients": sorted(lista_resultados, key=lambda ingredient: ingredient["precio"])[:10] 
    }


@router.get("/scraping", response_model=list[Url])
async def links():
    # URL de la página web que deseas raspar
    url = 'https://www.sedeco.cdmx.gob.mx/servicios/servicio/seguimiento-de-precios-de-la-canasta-basica'

    # Realiza una solicitud GET a la página web
    response = requests.get(url)

    # Crear una lista para almacenar los datos
    data_list = []

    # Verifica si la solicitud fue exitosa (código de estado 200)
    if response.status_code == 200:
        # Parsea el contenido HTML de la página web
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Encuentra el elemento con la clase 'Text'
        text_element = soup.find('div', class_='Text')  # Reemplaza 'Text' con la clase real

        # Verifica si se encontró el elemento con la clase 'Text'
        if text_element:
            
            # Busca todas las etiquetas <ul> dentro de ese elemento
            ul_tags = text_element.find_all('ul')

            # Itera a través de las etiquetas <ul> y busca las que contienen elementos <li>
            for ul_tag in ul_tags:
                # Verifica si la etiqueta <ul> tiene una etiqueta <ul> padre
                if not ul_tag.find_parent('ul'):
                    li_elements = ul_tag.find_all('li')

                    # Verifica si esta etiqueta <ul> contiene elementos <li>
                    if li_elements:
                        # Itera a través de los elementos <li>
                        for li in li_elements:
                            strong_tag = li.find('strong')  # Busca la etiqueta <strong> dentro del <li>
                            a_tag = li.find('a')  # Busca la etiqueta <a> dentro del <li>

                            if strong_tag:
                                strong_content = strong_tag.text.strip()  # Obtiene el contenido del <strong>
                                href = a_tag.get('href') if a_tag else ""  # Obtiene el atributo 'href' del <a> si existe
                                # Verifica si el contenido de <strong> no contiene "http" y termina con ".pdf"
                                if "http" not in strong_content and strong_content.endswith(".pdf"):
                                    data_list.append(Url(**{"name": strong_content, "url": "https://www.sedeco.cdmx.gob.mx"+href})
 )
    else:
        return {"error": "La solicitud GET no fue exitosa"}  

    return data_list

@router.get("/downloadDocuments", response_model=dict, status_code=status.HTTP_201_CREATED)
async def downloadDocument(): 

    documents = await links()
    # Obtiene la fecha actual
    now = datetime.now() 
    # Obtiene el año
    year = now.year 

    for document in documents:
        nombre = document.name
       
        if  str(year) in nombre.split(".")[0] and type(search_document("name",nombre.split(".")[0])) != Documento:
            # propaga la excepcion, no solo hace el return, te lanza el 404 y el json del detail
            # raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="El documento ya existe") 
    
            file_url = document.url
            response = requests.get(file_url, stream=True)
            
            # Divide la URL en partes usando "/" como separador y toma la última parte
            partes = file_url.split("/")
            nombre_pdf = partes[-1]
            

            # Nombre de las carpetas
            carpeta_static = "static"
            carpeta_docs = "docs"
            carpeta_pdf = "pdf"

            # Verificar si la carpeta "static" existe, si no, crearla
            if not os.path.exists(carpeta_static):
                os.makedirs(carpeta_static)

            # Verificar si la carpeta "docs" existe dentro de "static", si no, crearla
            carpeta_docs_path = os.path.join(carpeta_static, carpeta_docs)
            if not os.path.exists(carpeta_docs_path):
                os.makedirs(carpeta_docs_path)

            # Verificar si la carpeta "pdf" existe dentro de "docs", si no, crearla
            carpeta_pdf_path = os.path.join(carpeta_docs_path, carpeta_pdf)
            if not os.path.exists(carpeta_pdf_path):
                os.makedirs(carpeta_pdf_path)

            # Nombre completo del archivo PDF, incluyendo la ruta de carpetas
            nombre_pdf = os.path.join(carpeta_pdf_path, nombre_pdf)

            try:
                if response.status_code == 200:
                    with open(nombre_pdf, "wb") as pdf:
                        for chunk in response.iter_content(chunk_size=1024):
                            if chunk:
                                pdf.write(chunk) 
                
                    document_dict = {
                        "name":nombre.split(".")[0],
                        "name_document": convertPdfToWord(nombre_pdf),
                        "date": format_date(nombre.split(".")[0])
                    }
                    
                    id = db_client.documentos.insert_one(document_dict).inserted_id

                    # new_document = document_schema(db_client.documentos.find_one({"_id":id}) )

                    # return Documento(**new_document)

                else:
                    return {
                        "error": "Fallo en la descarga1"
                    }  
            except Exception as e:
                # Captura cualquier excepción y obtiene información detallada
                print("Se produjo una excepción:", str(e))
                return {"error": "Fallo en la descarga2"} 
    
    return {
         "documents": sorted(documents_schema(db_client.documentos.find()), key=lambda x: x["date"], reverse=True) 
     }

def format_date(nombre):
    # Dividir la cadena en palabras
    partes = nombre.split(" ")

    # Verificar si hay al menos tres partes (día, mes y año)
    if len(partes) >= 3:
        dia, mes_str, anio = partes[0],partes[2], partes[4]
           
    else:
        dia, mes_str, anio = 1, partes[0], partes[1]
    
    # Mapeo de nombres de meses a números de mes (en español)
    meses = {
        "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
        "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12
    }

    # Convertir el nombre del mes a número de mes
    mes_numero = meses.get(mes_str.lower(), 1)  # Si no se encuentra el mes, se usa enero por defecto

    # Convertir las partes de la fecha en enteros
    dia = int(dia)
    anio = int(anio)

    # Convertir la fecha a un objeto datetime
    fecha_datetime = datetime(anio, mes_numero, dia)

    # Formatear la fecha
    fecha_formateada = fecha_datetime.strftime("%Y-%m-%d")

    return fecha_formateada

# Definir una función para extraer y uniformar la fecha
def extract_and_format_date(document):
    date_str = document 
    # Intentar analizar la fecha en varios formatos
    try:
        date = datetime.strptime(date_str, "%d de %B de %Y")
    except ValueError:
        try:
            date = datetime.strptime(date_str, "%B %Y")
        except ValueError:
            date = datetime.min  # Establecer una fecha mínima para manejar casos no válidos
    return date

@router.get("/list-documents", response_model=dict)
async def documentos():
     return {
         "documents": sorted(documents_schema(db_client.documentos.find()), key=lambda x: x["date"], reverse=True) 
     }

# crear funcion, key generico por el objectId
def search_document( field: str,  key ):
    
    try:
        document = db_client.documentos.find_one({field: key}) 
     
     
        return Documento(**document_schema(document))
    except:
        return {"error":"No se ah encontrado el usuario"}
    
 
def convertPdfToWord(pdf_path: str):
    # pdf_path = Ruta del archivo PDF de entrada  

    # Nombre de las carpetas
    carpeta_static = "static"
    carpeta_docs = "docs"
    carpeta_pdf = "word"

    # Verificar si la carpeta "static" existe, si no, crearla
    if not os.path.exists(carpeta_static):
        os.makedirs(carpeta_static)

    # Verificar si la carpeta "docs" existe dentro de "static", si no, crearla
    carpeta_docs_path = os.path.join(carpeta_static, carpeta_docs)
    if not os.path.exists(carpeta_docs_path):
        os.makedirs(carpeta_docs_path)

    # Verificar si la carpeta "pdf" existe dentro de "docs", si no, crearla
    carpeta_pdf_path = os.path.join(carpeta_docs_path, carpeta_pdf)
    if not os.path.exists(carpeta_pdf_path):
        os.makedirs(carpeta_pdf_path)

    nombre_completo_pdf = os.path.basename(pdf_path)

    # Ruta del archivo Word de salida
    docx_path = "output-" + nombre_completo_pdf.replace(".pdf", "") + ".docx"


    # Nombre completo del archivo PDF, incluyendo la ruta de carpetas
    docx_path = os.path.join(carpeta_pdf_path, docx_path)


    # Crear un convertidor
    cv = Converter(pdf_path)

    # Convertir el PDF a Word
    cv.convert(docx_path, start=0, end=None)

    # Cerrar el convertidor
    cv.close()
    return docx_path
#     print("Pdf convertido a word: "+docx_path)




def getDataframeToWordTables(docx_path: str):
    # docx_path: Ruta del archivo Word de entrada 

    # Cargar el documento
    doc = Document(docx_path)

    # Crear una lista para almacenar los DataFrames de las tablas
    tables_dfs = []

    # Crear una variable para almacenar el texto previo a cada tabla
    prev_text = ''

    # Iniciar un contador de tablas
    table_count = 0

    # Recorrer los elementos del cuerpo del documento
    for element in doc.element.body:
        # Verificar si el elemento es una tabla
        if element.tag.endswith('tbl'):
            table = doc.tables[table_count]
            table_count += 1
            data = []
            for row in table.rows:
                row_data = [prev_text]  # Agregar el texto previo a la fila como primer elemento
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                data.append(row_data)
            # Crear un DataFrame de Pandas con los datos de la tabla
            df = pd.DataFrame(data)
            tables_dfs.append(df)
        # Verificar si el elemento es un párrafo
        elif element.tag.endswith('p'):
            paragraph = element.xpath('.//w:t')
            if paragraph:
                # Almacenar el texto del párrafo como el texto previo a la siguiente tabla
                text = paragraph[0].text.strip()
                # Tomar las últimas 3 palabras del párrafo y agregar "de 2023"
                prev_text = ' '.join(text.split()[-3:]) + ' de 2023'

    # Combinar todos los DataFrames en uno solo
    combined_df = pd.concat(tables_dfs, ignore_index=True)

    # Devolver el DataFrame combinado
    return combined_df

async def proveedores():  
    
    lastDocument = await documentos()
  
    # # Obtén las columnas desde la segunda hasta la penúltima
    columnas_deseadas = getDataframeToWordTables(lastDocument["documents"][-1]["name_document"]).iloc[0, 2:-1].tolist()

    # Imprime la lista resultante
    return columnas_deseadas

@router.get("/proveedores", response_model=dict)
async def proveedoresUnicos():  
    
    lastDocument = await documentos()
    
    # # Obtén las columnas desde la segunda hasta la penúltima
    columnas_deseadas = getDataframeToWordTables(lastDocument["documents"][-1]["name_document"]).iloc[0, 2:-1].tolist()
   
    proveedores = {
        "proveedores": set(columnas_deseadas)
    }
    # Imprime la lista resultante
    return proveedores

# Definir una función para extraer el valor numérico de una cadena de texto
def extrayendo_valor_numerico(text):
    # Buscar un patrón que coincida con un valor numérico, con o sin el símbolo de dólar
    match = re.search(r'(\$)?(\d+\.\d{2})', text)
    # Si se encuentra un valor numérico, devolverlo como un número de punto flotante
    if match:
        return float(match.group(2))
    # Si no se encuentra un valor numérico, devolver NaN
    return float('nan')
 
    import dateparser

# Función para convertir fechas en español al formato deseado
def conversion_fecha(date_str):
    if isinstance(date_str, str):  # Verificar si la entrada es una cadena
        date_obj = dateparser.parse(date_str)
        if date_obj:
            return date_obj.strftime('%Y-%m-%d')
    return date_str

async def preparationData(docx_path: str):
    
    combined_df = getDataframeToWordTables(docx_path)
    
    # Creamos una lista para almacenar los índices de las filas que queremos eliminar
    rows_to_drop = []
    
    # Recorremos las filas del DataFrame
    for i, row in combined_df.iterrows():
        # Verificamos si alguna celda de la fila contiene '+ Bajo' o '+ Alto'
        if any('+ Bajo' in str(cell) or '+ Alto' in str(cell) for cell in row):
            rows_to_drop.append(i)  # Agregamos el índice de la fila a la lista de filas a eliminar
            # Verificamos si la fila anterior tiene la palabra 'Tienda' o 'Mercado'
            if i > 0 and any('Tienda' in str(cell) or 'Mercado' in str(cell) for cell in combined_df.iloc[i-1]):
                rows_to_drop.append(i-1)  # Agregamos el índice de la fila anterior a la lista de filas a eliminar

    # Eliminamos las filas que cumplen con las condiciones
    combined_df.drop(rows_to_drop, inplace=True)

    # Reiniciamos el índice del DataFrame
    combined_df.reset_index(drop=True, inplace=True)
    
    # Obtener el nombre de la última columna
    last_column = combined_df.columns[-1]

    # Eliminar la última columna
    combined_df.drop(columns=[last_column], inplace=True)
    
    # Aplicar la función a cada celda de la tercera columna en adelante
    for column in combined_df.columns[2:]:
        combined_df[column] = combined_df[column].apply(extrayendo_valor_numerico)
    
    # Supongamos que tienes un DataFrame llamado 'combined_df' con 10 columnas
    nombres_columnas = ['fecha', 'ingrediente']

    # Cambia el nombre de las dos primeras columnas
    combined_df.rename(columns={combined_df.columns[0]: 'fecha', combined_df.columns[1]: 'ingrediente'}, inplace=True)

    # Agrega los nombres de columnas restantes desde la lista
    nombres_columnas.extend(await proveedores())

    # Asigna los nombres de columnas al DataFrame
    combined_df.columns = nombres_columnas
     

    return combined_df

 
async def getPriceProveedor(proveedor: str, precio: str, docx_path: str):
     # Obtener el DataFrame desde preparationData 
    df = await preparationData(docx_path)
    
    # Obtener las dos primeras columnas
    columnas_primeras_dos = df[['fecha', 'ingrediente']]
    
    # Obtener las columnas que coinciden con el nombre del proveedor
    columnas_proveedor = df.filter(like=proveedor)
    
    # Concatenar las dos partes en un nuevo DataFrame
    nuevo_df = pd.concat([columnas_primeras_dos, columnas_proveedor], axis=1)
    
    # Seleccionar la columna repetida basada en el valor de 'precio'
    if precio == 0:
        columna_seleccionada = columnas_proveedor.iloc[:, 0]  # Primera columna repetida
    else:
        columna_seleccionada = columnas_proveedor.iloc[:, 1]  # Segunda columna repetida
    
    # Concatenar las dos partes en un nuevo DataFrame
    nuevo_df = pd.concat([columnas_primeras_dos, columna_seleccionada], axis=1)
    
    # Imprimir el nuevo DataFrame o hacer lo que necesites con él
    
    # Cambiar el nombre de la columna 'promedio' a 'precio'
    nuevo_df.rename(columns={proveedor: 'precio'}, inplace=True)

    nuevo_df['fecha'] = nuevo_df['fecha'].astype(str)

    # Aplicar la función de conversión a la columna de fecha
    nuevo_df['fecha'] = nuevo_df['fecha'].apply(conversion_fecha)
    
    return nuevo_df
    

def arima_predictor(df, insumo, p, d, q):
    # Filtra el dataframe para obtener los datos del producto "Producto1"
    df_insumo = df[(df["ingrediente"] == insumo) & (df["precio"] != 0)][["fecha", "precio"]]
    # Convierte la columna "Fecha" a formato de fecha
    df_insumo["fecha"] = pd.to_datetime(df_insumo["fecha"])
    # Elimina las fechas duplicadas
    df_insumo = df_insumo.drop_duplicates(subset=['fecha'])
    # Establece la columna "Fecha" como el índice del DataFrame
    df_insumo = df_insumo.set_index("fecha")
    # Especifica la frecuencia de las fechas en el índice
    df_insumo = df_insumo.asfreq(freq='D')
    # Crea modelo ARIMA
    model = ARIMA(df_insumo, order=(p, d, q))
    # Ajusta el modelo a los datos
    model_fit = model.fit()
    # Predice precios para los próximos 7 días
    predicciones = model_fit.forecast(steps=7)
    # Imprime las predicciones de los 7 días próximos
    return df_insumo, predicciones


plt.rcParams["figure.figsize"] = (16, 5) 
def plot_predictions(df_insumo, predicciones, insumo):
    # Graficar los datos históricos
    plt.plot(df_insumo, label='Datos históricos')
    # Graficar las predicciones
    plt.plot(predicciones, label='Predicciones')
    # Agregar título y etiquetas a los ejes
    plt.title('Predicciones de precios para ' + insumo)
    plt.xlabel('Fecha')
    plt.ylabel('Precio')
    # Mostrar leyenda y graficar
    plt.legend()
    plt.show()